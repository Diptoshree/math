import streamlit as st
from docx import Document
import google.generativeai as genai

# -------------------- Setup Gemini --------------------
genai.configure(api_key="AIzaSyAymA7P4sRcrb0S9IAsC2cgkfxc0a6Fzj8")
model = genai.GenerativeModel("gemini-1.5-flash")

# -------------------- Helper Functions --------------------
def solve_math_question(question: str) -> str:
    prompt = f"""
You are a math tutor for high school students. Solve the following math problem step-by-step, 
clearly explaining each step in simple language that a student can understand.

Problem: {question}

Your answer:
"""
    response = model.generate_content(prompt)
    return response.text.strip()

def create_docx(question: str, answer: str) -> bytes:
    doc = Document()
    doc.add_heading("Math Problem Solver", level=1)
    doc.add_paragraph("📘 Question:")
    doc.add_paragraph(question)
    doc.add_paragraph("🧠 Step-by-step Answer:")
    doc.add_paragraph(answer)

    # Save to a BytesIO stream
    from io import BytesIO
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# -------------------- Streamlit UI --------------------
st.set_page_config(page_title="Math Problem Solver", page_icon="🧮")

st.title("🧮 Math Problem Solver for Students")
st.markdown("Enter any math question and get a clear step-by-step explanation.")

question = st.text_area("Type your math question here:")

if st.button("🔍 Solve"):
    if question.strip() == "":
        st.warning("Please enter a math question.")
    else:
        with st.spinner("Thinking..."):
            answer = solve_math_question(question)
        st.success("Here's your step-by-step solution:")
        st.markdown(f"**📘 Question:** {question}")
        st.markdown("**🧠 Answer:**")
        st.markdown(answer)

        # Generate Word Document
        docx_file = create_docx(question, answer)

        st.download_button(
            label="📥 Download as Word Document",
            data=docx_file,
            file_name="math_solution.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
